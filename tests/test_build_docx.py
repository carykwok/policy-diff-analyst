from docx import Document as DocxDocument
from scripts.build_docx import build_report

def test_build_report_writes_docx_with_headings(tmp_path):
    sections = [
        ("摘要", ["本文对比 2023 和 2024 政府工作报告。", "核心结论：GDP 目标下调，政策重点转向新质生产力。"]),
        ("结论", ["成长 > 价值；内需 > 外需；进攻 > 防守。"]),
    ]
    disclaimer = "本文仅供研究参考，不构成投资建议。"
    out = tmp_path / "report.docx"
    build_report(
        title="2024 vs 2023 政府工作报告分析",
        style="research",
        sections=sections,
        disclaimer=disclaimer,
        output_path=out,
    )
    assert out.exists()
    doc = DocxDocument(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "摘要" in headings
    assert "结论" in headings
    full = "\n".join(p.text for p in doc.paragraphs)
    assert disclaimer in full
