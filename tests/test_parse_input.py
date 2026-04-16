from pathlib import Path
from docx import Document as DocxDocument
from pypdf import PdfWriter
import pytest
from scripts.parse_input import parse_text, parse_docx, parse_pdf, detect_sections

def test_parse_text_extracts_title_and_sections(fixtures_dir):
    raw = (fixtures_dir / "sample_report.txt").read_text(encoding="utf-8")
    doc = parse_text(raw, year=2024, title="2024 政府工作报告")
    assert doc.year == 2024
    assert len(doc.sections) == 2
    assert doc.sections[0].heading.startswith("一、")
    assert "126 万亿元" in doc.sections[0].body

def test_detect_sections_numbered_chinese():
    body = "一、回顾\n去年 GDP 增长 5.2%\n二、展望\n今年目标 5% 左右"
    sections = detect_sections(body)
    assert len(sections) == 2
    assert sections[0].heading == "一、回顾"
    assert sections[1].body.startswith("今年目标")

def test_parse_docx_reads_paragraphs(tmp_path):
    docx_path = tmp_path / "sample.docx"
    d = DocxDocument()
    d.add_paragraph("2024 政府工作报告")
    d.add_paragraph("一、回顾")
    d.add_paragraph("GDP 5.2%")
    d.save(docx_path)
    doc = parse_docx(docx_path, year=2024)
    assert "GDP 5.2%" in doc.raw_text
    assert len(doc.sections) == 1

def test_parse_pdf_extracts_text(tmp_path):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    pdf_path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "2024 政府工作报告")
    c.drawString(100, 730, "一、回顾")
    c.drawString(100, 710, "GDP 5.2%")
    c.save()
    doc = parse_pdf(pdf_path, year=2024)
    assert "GDP" in doc.raw_text
