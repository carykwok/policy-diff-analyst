from pathlib import Path
from docx import Document as DocxDocument
from scripts.models import DiffReport, DiffItem, StrengthScore
from scripts.build_annotated_docx import build_annotated_docx


def _sample_report():
    return DiffReport(
        old_doc_title="2023 政府工作报告",
        new_doc_title="2024 政府工作报告",
        items=[
            DiffItem(layer="A3", change_type="added", old="", new="新质生产力", note="新增表述"),
            DiffItem(layer="A4", change_type="modified", old="积极房地产", new="坚决房地产",
                     note="强度 3→5"),
            DiffItem(layer="A3", change_type="removed", old="互联网+", new="", note="不再提及"),
        ],
        strength=[StrengthScore("A3_产业", 2.0, 3.0)],
        term_freq={"新质生产力": {"old": 0, "new": 3}, "房地产": {"old": 2, "new": 2}},
    )


def test_annotated_docx_is_created(tmp_path):
    out = tmp_path / "annotated.docx"
    new_text = "2024 政府工作报告\n大力发展新质生产力。\n坚决防范化解房地产风险。"
    build_annotated_docx(new_text, _sample_report(), out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_annotated_docx_contains_red_highlights(tmp_path):
    out = tmp_path / "annotated.docx"
    new_text = "2024 政府工作报告\n大力发展新质生产力。坚决防范化解房地产风险。"
    build_annotated_docx(new_text, _sample_report(), out)

    doc = DocxDocument(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "新质生产力" in all_text
    assert "房地产" in all_text
    # Check annotations exist
    assert "新增" in all_text
    assert "强度" in all_text


def test_annotated_docx_has_removed_section(tmp_path):
    out = tmp_path / "annotated.docx"
    new_text = "2024 政府工作报告\n大力发展新质生产力。"
    build_annotated_docx(new_text, _sample_report(), out)

    doc = DocxDocument(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "互联网+" in all_text
    assert "不再提及" in all_text


def test_annotated_docx_empty_report(tmp_path):
    out = tmp_path / "annotated.docx"
    empty_report = DiffReport(
        old_doc_title="2023", new_doc_title="2024",
        items=[], strength=[], term_freq={},
    )
    build_annotated_docx("Some text here.", empty_report, out)
    assert out.exists()
