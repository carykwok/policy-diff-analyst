from pathlib import Path
from docx import Document
from docx.shared import Pt

def build_report(
    *,
    title: str,
    style: str,
    sections: list[tuple[str, list[str]]],
    disclaimer: str,
    output_path: Path,
) -> None:
    doc = Document()

    # Title
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"风格：{style}").italic = True

    # Body
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for para in paragraphs:
            doc.add_paragraph(para)

    # Disclaimer at end
    doc.add_heading("免责声明", level=2)
    p = doc.add_paragraph(disclaimer)
    p.runs[0].font.size = Pt(9)

    doc.save(output_path)
