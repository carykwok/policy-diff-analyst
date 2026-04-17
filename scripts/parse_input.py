import re
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
from scripts.models import Document, Section

# Matches headings like "一、", "二、", "三、" at line start
_SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、.+", re.MULTILINE)

def detect_sections(body: str) -> list[Section]:
    matches = list(_SECTION_HEADING_RE.finditer(body))
    if not matches:
        return [Section(heading="全文", body=body.strip())]
    sections: list[Section] = []
    for i, m in enumerate(matches):
        heading_line = m.group(0).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(body)
        content = body[start:end].strip()
        sections.append(Section(heading=heading_line, body=content))
    return sections

def parse_text(raw: str, *, year: int, file_type: str, title: str | None = None, source_url: str | None = None) -> Document:
    lines = raw.strip().splitlines()
    inferred_title = title or (lines[0].strip() if lines else "")
    body = "\n".join(lines[1:]) if title is None and lines else raw
    sections = detect_sections(body)
    return Document(
        title=inferred_title,
        year=year,
        file_type=file_type,
        source_url=source_url,
        sections=sections,
        raw_text=raw,
    )

def parse_docx(path: Path, *, year: int, file_type: str, source_url: str | None = None) -> Document:
    d = DocxDocument(str(path))
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    raw = "\n".join(paras)
    return parse_text(raw, year=year, file_type=file_type, title=paras[0] if paras else "", source_url=source_url)

def parse_pdf(path: Path, *, year: int, file_type: str, source_url: str | None = None) -> Document:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    raw = "\n".join(pages)
    return parse_text(raw, year=year, file_type=file_type, source_url=source_url)
