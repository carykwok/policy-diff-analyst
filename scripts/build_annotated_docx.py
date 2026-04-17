"""Build an annotated copy of the new document with red highlights and grey notes.

For each keyword in the DiffReport:
  - added keywords: marked in red bold
  - removed keywords: noted at end of paragraph as grey text
  - modified (strength change): marked in red with grey strength annotation

Output: a .docx where the user can see exactly what changed in-situ.
"""
from __future__ import annotations

import re
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from scripts.models import DiffReport, DiffItem, StructureDiff

_RED = RGBColor(0xCC, 0x00, 0x00)
_GREY = RGBColor(0x88, 0x88, 0x88)


def build_annotated_docx(
    new_text: str,
    report: DiffReport,
    output_path: Path,
    *,
    structure_diff: StructureDiff | None = None,
) -> None:
    doc = DocxDocument()

    # Title
    doc.add_heading(f"{report.new_doc_title}（标注版）", level=0)
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"对比基准：{report.old_doc_title} → {report.new_doc_title}")
    run.italic = True
    run.font.color.rgb = _GREY

    # ── Structure outline summary (framework-level view first) ──
    if structure_diff:
        doc.add_heading("一、框架结构对比", level=1)
        summary_p = doc.add_paragraph()
        summary_run = summary_p.add_run(structure_diff.summary)
        summary_run.bold = True

        for change in structure_diff.changes:
            p = doc.add_paragraph(style="List Bullet")
            if change.change_type == "added":
                r = p.add_run(f"【新增】{change.heading_new}")
                r.font.color.rgb = _RED
                r.bold = True
            elif change.change_type == "removed":
                r = p.add_run(f"【删除】{change.heading_old}")
                r.font.color.rgb = _GREY
            elif change.change_type in ("moved_up", "moved_down"):
                r = p.add_run(f"{change.heading_new}")
                r.bold = True
                note_r = p.add_run(f"  （{change.note}）")
                note_r.font.color.rgb = _GREY
                note_r.font.size = Pt(8)
            elif change.change_type == "renamed":
                r = p.add_run(f"{change.heading_new}")
                r.font.color.rgb = _RED
                note_r = p.add_run(f"  （{change.note}）")
                note_r.font.color.rgb = _GREY
                note_r.font.size = Pt(8)
            else:  # kept
                p.add_run(change.heading_new)

        doc.add_paragraph()  # spacer

    doc.add_heading("二、逐段标注" if structure_diff else "逐段标注", level=1)

    # Index diff items for quick lookup
    added_kws: dict[str, DiffItem] = {}
    modified_kws: dict[str, DiffItem] = {}
    removed_items: list[DiffItem] = []
    # term_freq keys are bare keywords (no modifier prefix)
    bare_keywords = set(report.term_freq.keys())
    for item in report.items:
        if item.change_type == "added" and item.new:
            added_kws[item.new] = item
        elif item.change_type == "modified" and (item.new or item.old):
            # DiffItem.new may be "坚决房地产" (modifier+keyword).
            # Find the bare keyword by checking which term_freq key is contained.
            candidate = item.new or item.old
            matched = False
            for kw in bare_keywords:
                if kw in candidate:
                    modified_kws[kw] = item
                    matched = True
                    break
            if not matched:
                modified_kws[candidate] = item
        elif item.change_type == "removed" and item.old:
            removed_items.append(item)

    # Build a combined regex for all keywords to highlight
    highlight_map: dict[str, str] = {}  # keyword -> "added" | "modified"
    for kw in added_kws:
        highlight_map[kw] = "added"
    for kw in modified_kws:
        highlight_map[kw] = "modified"

    # Process the new text paragraph by paragraph
    paragraphs = new_text.strip().split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        p = doc.add_paragraph()

        if not highlight_map:
            p.add_run(para_text)
            continue

        # Sort keywords by length (longest first) to avoid partial matches
        sorted_kws = sorted(highlight_map.keys(), key=len, reverse=True)
        # Build pattern matching any keyword
        pattern = "|".join(re.escape(kw) for kw in sorted_kws if kw in para_text)

        if not pattern:
            p.add_run(para_text)
            _append_removed_notes(p, para_text, removed_items)
            continue

        # Split text by keyword matches, preserving matched segments
        last_end = 0
        for match in re.finditer(pattern, para_text):
            # Plain text before the match
            if match.start() > last_end:
                p.add_run(para_text[last_end:match.start()])

            kw = match.group()
            change_type = highlight_map.get(kw, "modified")
            item = added_kws.get(kw) or modified_kws.get(kw)

            # Red highlighted keyword
            kw_run = p.add_run(kw)
            kw_run.font.color.rgb = _RED
            kw_run.bold = True

            # Grey inline annotation
            if item:
                if change_type == "added":
                    note_text = f"【新增·{item.layer}】"
                else:
                    note_text = f"【{item.note}】"
                note_run = p.add_run(note_text)
                note_run.font.color.rgb = _GREY
                note_run.font.size = Pt(8)

            last_end = match.end()

        # Remaining text after last match
        if last_end < len(para_text):
            p.add_run(para_text[last_end:])

        _append_removed_notes(p, para_text, removed_items)

    # Append a section for all removed keywords at the end
    if removed_items:
        doc.add_heading("本版不再提及的表述", level=2)
        for item in removed_items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{item.old}")
            run.font.color.rgb = _GREY
            note = p.add_run(f"  （{item.layer}·{item.note}）")
            note.font.color.rgb = _GREY
            note.font.size = Pt(8)

    doc.save(output_path)


def _append_removed_notes(p, para_text: str, removed_items: list[DiffItem]) -> None:
    """If any removed keyword was in the OLD text and this paragraph is about
    the same topic (same layer keywords nearby), add a grey note."""
    # We don't try to match removed keywords to specific paragraphs —
    # they're collected at the end of the document instead.
    pass
